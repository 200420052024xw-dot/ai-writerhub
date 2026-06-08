import json
import os
from io import BytesIO

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from fastapi import HTTPException

from app.prompts.format import FORMAT_STRING_FIELDS, FORMAT_BOOL_FIELDS, FONT_OPTIONS, FONT_SIZE_OPTIONS, build_format_organize_prompts, build_format_parse_prompts
from app.schemas.format import FormatConfig, FormatDocumentParagraph, FormatExportDocxRequest, FormatOrganizeResponse
from app.services.llm_client import RuntimeModelConfig, call_chat_model


FONT_SIZE_PT = {
    "八号": 5,
    "七号": 5.5,
    "小六": 6.5,
    "六号": 7.5,
    "小五": 9,
    "五号": 10.5,
    "小四": 12,
    "四号": 14,
    "小三": 15,
    "三号": 16,
    "小二": 18,
    "二号": 22,
    "小一": 24,
    "一号": 26,
    "小初": 36,
    "初号": 42,
}

# Normalization maps for each option type
_FONT_NORM = {
    "宋体": "宋体（SimSun）", "simsun": "宋体（SimSun）", "宋体（simsun）": "宋体（SimSun）",
    "微软雅黑": "微软雅黑", "雅黑": "微软雅黑",
    "黑体": "黑体", "仿宋": "仿宋", "楷体": "楷体",
}

_FONTSIZE_NORM = {
    "初号": "初号（42pt）", "42pt": "初号（42pt）", "42": "初号（42pt）",
    "小初": "小初（36pt）", "36pt": "小初（36pt）", "36": "小初（36pt）",
    "一号": "一号（26pt）", "26pt": "一号（26pt）", "26": "一号（26pt）",
    "小一": "小一（24pt）", "24pt": "小一（24pt）", "24": "小一（24pt）",
    "二号": "二号（22pt）", "22pt": "二号（22pt）", "22": "二号（22pt）",
    "小二": "小二（18pt）", "18pt": "小二（18pt）", "18": "小二（18pt）",
    "三号": "三号（16pt）", "16pt": "三号（16pt）", "16": "三号（16pt）",
    "小三": "小三（15pt）", "15pt": "小三（15pt）", "15": "小三（15pt）",
    "四号": "四号（14pt）", "14pt": "四号（14pt）", "14": "四号（14pt）",
    "小四": "小四（12pt）", "12pt": "小四（12pt）", "12": "小四（12pt）",
    "五号": "五号（10.5pt）", "10.5pt": "五号（10.5pt）", "10.5": "五号（10.5pt）",
    "小五": "小五（9pt）", "9pt": "小五（9pt）", "9": "小五（9pt）",
    "六号": "六号（7.5pt）", "7.5pt": "六号（7.5pt）", "7.5": "六号（7.5pt）",
    "小六": "小六（6.5pt）", "6.5pt": "小六（6.5pt）", "6.5": "小六（6.5pt）",
    "七号": "七号（5.5pt）", "5.5pt": "七号（5.5pt）", "5.5": "七号（5.5pt）",
    "八号": "八号（5pt）", "5pt": "八号（5pt）", "5": "八号（5pt）",
}

_LINEHEIGHT_NORM = {
    "单倍": "单倍行距", "1": "单倍行距", "1.0": "单倍行距",
    "1.25": "1.25 倍行距", "1.5": "1.5 倍行距",
    "双倍": "2 倍行距", "2": "2 倍行距", "2.0": "2 倍行距",
}

_INDENT_NORM = {
    "无": "无缩进", "不缩进": "无缩进",
    "首行": "首行缩进 2 字符", "首行缩进": "首行缩进 2 字符", "2字符": "首行缩进 2 字符",
    "左缩进": "左缩进 2 字符", "悬挂": "悬挂缩进 2 字符", "悬挂缩进": "悬挂缩进 2 字符",
}

_ALIGN_NORM = {
    "左": "左对齐", "居中": "居中对齐", "中间": "居中对齐",
    "右": "右对齐", "两端": "两端对齐", "分散": "两端对齐",
}

_PAPERSIZE_NORM = {
    "a4": "A4（21 × 29.7cm）", "A4": "A4（21 × 29.7cm）",
    "a5": "A5（14.8 × 21cm）", "A5": "A5（14.8 × 21cm）",
    "b5": "B5（17.6 × 25cm）", "B5": "B5（17.6 × 25cm）",
    "letter": "Letter（21.6 × 27.9cm）", "Letter": "Letter（21.6 × 27.9cm）",
}

_MARGIN_NORM = {
    "普通": "普通：上/下 2.54cm，左/右 3.18cm", "标准": "普通：上/下 2.54cm，左/右 3.18cm",
    "默认": "普通：上/下 2.54cm，左/右 3.18cm",
    "窄": "窄边距：上/下/左/右 1.27cm", "窄边距": "窄边距：上/下/左/右 1.27cm",
    "小": "窄边距：上/下/左/右 1.27cm",
}

# Map each config field to its normalization table
_FIELD_NORM: dict[str, dict[str, str]] = {}
for _f in ["bodyFont", "titleFont", "h1Font", "h2Font", "h3Font"]:
    _FIELD_NORM[_f] = _FONT_NORM
for _f in ["bodyFontSize", "titleFontSize", "h1FontSize", "h2FontSize", "h3FontSize"]:
    _FIELD_NORM[_f] = _FONTSIZE_NORM
_FIELD_NORM["lineHeight"] = _LINEHEIGHT_NORM
_FIELD_NORM["indent"] = _INDENT_NORM
_FIELD_NORM["align"] = _ALIGN_NORM
_FIELD_NORM["paperSize"] = _PAPERSIZE_NORM
_FIELD_NORM["orientation"] = {"纵向": "纵向", "横向": "横向", "landscape": "横向", "portrait": "纵向"}
_FIELD_NORM["margin"] = _MARGIN_NORM


def _normalize_field(key: str, value: str) -> str:
    """Map LLM free-form value to the closest UI option string."""
    norm_table = _FIELD_NORM.get(key)
    if not norm_table:
        return value

    # Direct match (case-insensitive)
    value_lower = value.strip().lower()
    for k, v in norm_table.items():
        if k.lower() == value_lower:
            return v

    # Substring match
    for k, v in norm_table.items():
        if k.lower() in value_lower or value_lower in k.lower():
            return v

    return value


async def parse_format_config(prompt: str, current_config: FormatConfig, model_config: RuntimeModelConfig) -> FormatConfig:
    system_prompt, user_prompt = build_format_parse_prompts(prompt, current_config.model_dump())
    raw = await call_chat_model(system_prompt, user_prompt, model_config)
    content = raw.strip()
    if content.startswith("```"):
        content = content.strip("`").removeprefix("json").strip()
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail="模型未返回有效 JSON") from exc

    current = current_config.model_dump()
    extra_parts: list[str] = []
    existing_extra = current.get("extraRequirements", "")
    if existing_extra:
        extra_parts.append(existing_extra)

    # Process string fields
    for key in FORMAT_STRING_FIELDS:
        value = parsed.get(key)
        if not isinstance(value, str) or not value.strip():
            continue

        if key == "extraRequirements":
            extra_parts.append(value.strip())
        elif key in ("header", "footer"):
            current[key] = value.strip()
        else:
            normalized = _normalize_field(key, value.strip())
            if key in _FIELD_NORM:
                # Has a normalization table — only accept if normalized value is a known option
                norm_table = _FIELD_NORM[key]
                if normalized in norm_table.values():
                    current[key] = normalized
                else:
                    extra_parts.append(f"{key}: {value.strip()}")
            else:
                current[key] = normalized

    # Process boolean fields
    for key in FORMAT_BOOL_FIELDS:
        value = parsed.get(key)
        if isinstance(value, bool):
            current[key] = value

    if extra_parts:
        current["extraRequirements"] = "；".join(extra_parts)

    return FormatConfig(**current)


def normalize_east_asia_font(value: str) -> str:
    """Normalize frontend font name to standard Chinese font name for OOXML w:eastAsia."""
    if not value:
        return "宋体"

    value = value.strip()

    if "宋体" in value or "SimSun" in value or "simsun" in value:
        return "宋体"
    if "微软雅黑" in value or "雅黑" in value:
        return "微软雅黑"
    if "黑体" in value:
        return "黑体"
    if "仿宋" in value:
        return "仿宋"
    if "楷体" in value:
        return "楷体"

    return "宋体"


def parse_font_size(value: str) -> float:
    for label, size in FONT_SIZE_PT.items():
        if label in value:
            return size
    for token in value.replace("pt", "").replace("（", " ").replace("）", " ").replace("(", " ").replace(")", " ").split():
        try:
            return float(token)
        except ValueError:
            continue
    return 12


def parse_line_spacing(value: str) -> float:
    if "2" in value:
        return 2
    if "1.25" in value:
        return 1.25
    if "1.5" in value:
        return 1.5
    return 1


def paragraph_alignment(value: str):
    if "居中" in value:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "右" in value:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if "两端" in value:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def configure_section(document: Document, config: FormatConfig) -> None:
    section = document.sections[0]
    if "A5" in config.paperSize:
        w, h = Cm(14.8), Cm(21)
    elif "B5" in config.paperSize:
        w, h = Cm(17.6), Cm(25)
    elif "Letter" in config.paperSize:
        w, h = Cm(21.6), Cm(27.9)
    else:
        w, h = Cm(21), Cm(29.7)

    if "横" in (config.orientation or ""):
        section.page_width = h
        section.page_height = w
    else:
        section.page_width = w
        section.page_height = h

    if "窄" in config.margin or "窄" in config.extraRequirements:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.27)
    else:
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(3.18)

    if config.header.strip():
        section.header.paragraphs[0].text = config.header.strip()
    if config.footer.strip():
        section.footer.paragraphs[0].text = config.footer.strip()


def apply_run_font(run, east_asia_font: str, ascii_font: str, size_pt: float, bold: bool = False) -> None:
    """Set run font with full OOXML control for both Chinese and Western characters."""
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr.get_or_add_rFonts()

    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)


def apply_paragraph_style(paragraph, config: FormatConfig) -> None:
    paragraph.alignment = paragraph_alignment(config.align)
    paragraph.paragraph_format.line_spacing = parse_line_spacing(config.lineHeight)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    indent = config.indent or ""
    if "首行" in indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    elif "左缩进" in indent:
        paragraph.paragraph_format.left_indent = Cm(0.74)
    elif "悬挂" in indent:
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        paragraph.paragraph_format.left_indent = Cm(0.74)


def add_heading_paragraph(document: Document, block: FormatDocumentParagraph, config: FormatConfig) -> None:
    """Add a heading paragraph with manual formatting (no built-in heading style)."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(block.content or "")

    level = max(1, min(3, block.level or 1))

    # Default values by level
    defaults = {
        1: ("黑体", 16.0, True, WD_ALIGN_PARAGRAPH.CENTER),
        2: ("黑体", 14.0, True, WD_ALIGN_PARAGRAPH.LEFT),
        3: ("黑体", 12.0, True, WD_ALIGN_PARAGRAPH.LEFT),
    }
    default_font, default_size, default_bold, default_align = defaults[level]

    # Read per-level config
    font_val = getattr(config, f"h{level}Font", "") or ""
    size_val = getattr(config, f"h{level}FontSize", "") or ""
    bold_val = getattr(config, f"h{level}Bold", True)

    east_asia_font = normalize_east_asia_font(font_val) if font_val else default_font
    size_pt = parse_font_size(size_val) if size_val else default_size
    bold = bold_val

    paragraph.alignment = default_align
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)

    apply_run_font(run, east_asia_font=east_asia_font, ascii_font="Times New Roman", size_pt=size_pt, bold=bold)


def _apply_body_font(run, config: FormatConfig) -> None:
    """Apply body text font settings to a run."""
    east_asia_font = normalize_east_asia_font(config.bodyFont) if config.bodyFont else "宋体"
    size_pt = parse_font_size(config.bodyFontSize) if config.bodyFontSize else 12
    apply_run_font(run, east_asia_font, "Times New Roman", size_pt, bold=config.bodyBold)


def add_paragraph(document: Document, block: FormatDocumentParagraph, config: FormatConfig) -> None:
    if block.type == "title":
        return

    if block.type == "heading":
        add_heading_paragraph(document, block, config)
        return

    if block.type == "list":
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(block.content or "")
        apply_paragraph_style(paragraph, config)
        _apply_body_font(run, config)
        return

    # paragraph / table / default
    paragraph = document.add_paragraph()
    run = paragraph.add_run(block.content or "")
    apply_paragraph_style(paragraph, config)
    _apply_body_font(run, config)


async def organize_document(
    paragraphs: list[FormatDocumentParagraph],
    config: FormatConfig,
    model_config: RuntimeModelConfig,
) -> FormatOrganizeResponse:
    system_prompt, user_prompt = build_format_organize_prompts(
        [paragraph.model_dump() for paragraph in paragraphs],
        config.model_dump(),
    )
    raw = await call_chat_model(system_prompt, user_prompt, model_config)
    content = raw.strip()
    if content.startswith("```"):
        content = content.strip("`").removeprefix("json").strip()
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail="模型未返回有效 JSON") from exc

    paragraph_data = parsed.get("paragraphs", [])
    if not isinstance(paragraph_data, list):
        raise HTTPException(status_code=502, detail="Model response is missing the paragraphs array")
    try:
        organized = [FormatDocumentParagraph(**item) for item in paragraph_data]
    except (TypeError, ValueError) as exc:
        raise HTTPException(status_code=502, detail="Model returned invalid structured paragraphs") from exc

    expected_ids = [paragraph.paragraph_id for paragraph in paragraphs]
    returned_ids = [paragraph.paragraph_id for paragraph in organized]
    if returned_ids != expected_ids:
        raise HTTPException(status_code=502, detail="Model changed paragraph IDs or paragraph order")
    if len(set(returned_ids)) != len(returned_ids):
        raise HTTPException(status_code=502, detail="Model returned duplicate paragraph IDs")
    for paragraph in organized:
        if paragraph.type == "title" and paragraph.level != 0:
            raise HTTPException(status_code=502, detail="Model returned an invalid title level")
        if paragraph.type == "heading" and paragraph.level == 0:
            raise HTTPException(status_code=502, detail="Model returned an invalid heading level")
    return FormatOrganizeResponse(paragraphs=organized)


def build_docx(payload: FormatExportDocxRequest) -> BytesIO:
    document = Document()
    configure_section(document, payload.config)
    config = payload.config

    # Document title
    east_asia_font = normalize_east_asia_font(config.titleFont) if config.titleFont else "黑体"
    size_pt = parse_font_size(config.titleFontSize) if config.titleFontSize else 16
    bold = config.titleBold

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run(payload.title)
    apply_run_font(title_run, east_asia_font=east_asia_font, ascii_font="Times New Roman", size_pt=size_pt, bold=bold)

    for paragraph in payload.paragraphs:
        add_paragraph(document, paragraph, config)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


# --------------- PDF export (PyMuPDF) ---------------

_PDF_FONT_DIR = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts")
_PDF_FONT_MAP = {
    "宋体": "simsun.ttc", "simsun": "simsun.ttc",
    "黑体": "simhei.ttf", "simhei": "simhei.ttf",
    "微软雅黑": "msyh.ttc", "雅黑": "msyh.ttc",
    "仿宋": "simfang.ttf",
    "楷体": "simkai.ttf",
}


def _load_pdf_font(name: str) -> "fitz.Font":
    """Load a Chinese font from system fonts, fallback to built-in."""
    for key, filename in _PDF_FONT_MAP.items():
        if key in (name or ""):
            path = os.path.join(_PDF_FONT_DIR, filename)
            if os.path.exists(path):
                return fitz.Font(fontfile=path)
    # Fallback: built-in CJK font
    return fitz.Font("china-s")


def _pdf_paper_size(config: FormatConfig) -> tuple[float, float]:
    """Return (width, height) in points for the configured paper size."""
    sizes = {"A5": (419.53, 595.28), "B5": (498.90, 708.66), "Letter": (612.0, 792.0)}
    pw, ph = 595.28, 841.89  # A4 default
    for key, sz in sizes.items():
        if key in config.paperSize:
            pw, ph = sz
            break
    if "横" in (config.orientation or ""):
        pw, ph = ph, pw
    return pw, ph


def _pdf_margins(config: FormatConfig) -> tuple[float, float, float, float]:
    """Return (top, right, bottom, left) in points."""
    if "窄" in config.margin:
        return 36.0, 36.0, 36.0, 36.0
    return 72.0, 90.71, 72.0, 90.71


def _wrap_text(text: str, font: "fitz.Font", fontsize: float, max_width: float) -> list[str]:
    """Split text into lines that fit within *max_width* points."""
    if not text:
        return [""]
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if font.text_length(test, fontsize=fontsize) > max_width:
            if current:
                lines.append(current)
            current = ch
        else:
            current = test
    if current:
        lines.append(current)
    return lines or [""]


def build_pdf(title: str, paragraphs: list[FormatDocumentParagraph], config: FormatConfig) -> BytesIO:
    """Generate a PDF document using PyMuPDF."""
    pw, ph = _pdf_paper_size(config)
    mt, mr, mb, ml = _pdf_margins(config)
    content_w = pw - ml - mr

    body_size = parse_font_size(config.bodyFontSize) if config.bodyFontSize else 12.0
    line_spacing = parse_line_spacing(config.lineHeight)

    # Load fonts
    body_font = _load_pdf_font(config.bodyFont or "宋体")
    title_font = _load_pdf_font(config.titleFont or "黑体")

    heading_fonts: dict[int, "fitz.Font"] = {}
    heading_sizes: dict[int, float] = {}
    for level in (1, 2, 3):
        fval = getattr(config, f"h{level}Font", "") or ""
        sval = getattr(config, f"h{level}FontSize", "") or ""
        heading_fonts[level] = _load_pdf_font(fval) if fval else _load_pdf_font("黑体")
        defaults = {1: 16.0, 2: 14.0, 3: 12.0}
        heading_sizes[level] = parse_font_size(sval) if sval else defaults[level]

    # State
    doc = fitz.open()
    page = doc.new_page(width=pw, height=ph)
    y = mt
    registered: set[str] = set()

    def _register(font_obj: "fitz.Font", name: str) -> str:
        if name not in registered:
            page.insert_font(fontname=name, fontbuffer=font_obj.buffer)
            registered.add(name)
        return name

    def _new_page():
        nonlocal page, y, registered
        page = doc.new_page(width=pw, height=ph)
        y = mt
        registered = set()

    def _check_space(needed: float):
        nonlocal y
        if y + needed > ph - mb:
            _new_page()

    def _draw_text(text: str, font_obj: "fitz.Font", font_name: str,
                   fontsize: float, x_offset: float = 0, align: str = "left",
                   extra_gap: float = 0):
        nonlocal y
        if not text:
            return
        _register(font_obj, font_name)
        line_h = fontsize * line_spacing
        avail_w = content_w - x_offset
        lines = _wrap_text(text, font_obj, fontsize, avail_w)

        for i, line in enumerate(lines):
            _check_space(line_h)
            lw = font_obj.text_length(line, fontsize=fontsize)
            if align == "center":
                x = ml + (content_w - lw) / 2
            elif align == "right":
                x = ml + content_w - lw
            else:
                x = ml + (x_offset if i == 0 else 0)
            page.insert_text(fitz.Point(x, y + fontsize), line,
                             fontname=font_name, fontsize=fontsize)
            y += line_h
        if extra_gap:
            y += extra_gap

    # --- Title ---
    title_size = parse_font_size(config.titleFontSize) if config.titleFontSize else 16.0
    _draw_text(title, title_font, "pdf_title", title_size, align="center", extra_gap=title_size * 0.4)

    # --- Paragraphs ---
    for para in paragraphs:
        if para.type == "title":
            continue

        if para.type == "heading":
            level = max(1, min(3, para.level or 1))
            _draw_text(para.content, heading_fonts[level], f"pdf_h{level}",
                       heading_sizes[level],
                       align="center" if level == 1 else "left",
                       extra_gap=heading_sizes[level] * 0.25)

        elif para.type == "list":
            _draw_text(f"• {para.content}", body_font, "pdf_body", body_size, x_offset=18)

        else:  # paragraph / table / default
            indent = body_size * 2 if "首行" in (config.indent or "") else 0
            _draw_text(para.content, body_font, "pdf_body", body_size, x_offset=indent)

    buffer = BytesIO()
    doc.save(buffer)
    doc.close()
    buffer.seek(0)
    return buffer
